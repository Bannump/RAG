"""
Resume Builder Module - Build optimized resumes from experiences, skills, and portfolio
"""
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.my_personal_agent.core.rag_engine import RAGEngine
from src.my_personal_agent.utils.file_handler import FileHandler
from src.my_personal_agent.utils.text_processor import TextProcessor


class ResumeBuilder:
    """Build ATS-optimized resumes from user inputs"""
    
    def __init__(self, rag_engine: Optional[RAGEngine] = None):
        self.rag_engine = rag_engine or RAGEngine()
        self.file_handler = FileHandler()
        self.text_processor = TextProcessor()
    
    def build_resume(
        self,
        experiences: List[Dict[str, Any]],
        skills: List[str],
        education: List[Dict[str, Any]],
        portfolio_items: Optional[List[Dict[str, Any]]] = None,
        target_job: Optional[str] = None,
        output_format: str = "docx",
        user_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Build a resume from provided information
        
        Args:
            experiences: List of work experiences with keys: title, company, duration, description
            skills: List of skills
            education: List of education entries with keys: degree, institution, year, details
            portfolio_items: Optional list of portfolio projects
            target_job: Optional target job title/description for optimization
            output_format: "docx", "txt", or "pdf"
            user_id: Optional user ID for file organization
        
        Returns:
            Dictionary with resume content and file path
        """
        # Generate optimized content using LLM if target_job provided
        if target_job:
            optimized_content = self._optimize_for_job(
                experiences, skills, education, portfolio_items, target_job
            )
            experiences = optimized_content.get("experiences", experiences)
            skills = optimized_content.get("skills", skills)
        
        # Generate professional summary
        summary = self._generate_summary(experiences, skills, target_job)
        
        # Build resume sections
        resume_data = {
            "summary": summary,
            "experiences": experiences,
            "skills": skills,
            "education": education,
            "portfolio": portfolio_items or [],
        }
        
        # Create resume document
        if output_format == "docx":
            file_path = self._create_docx_resume(resume_data, user_id)
        elif output_format == "txt":
            file_path = self._create_txt_resume(resume_data, user_id)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
        
        # Generate resume text for analysis
        resume_text = self._generate_resume_text(resume_data)
        
        return {
            "resume_data": resume_data,
            "resume_text": resume_text,
            "file_path": file_path,
            "format": output_format,
        }
    
    def _optimize_for_job(
        self,
        experiences: List[Dict[str, Any]],
        skills: List[str],
        education: List[Dict[str, Any]],
        portfolio_items: Optional[List[Dict[str, Any]]],
        target_job: str,
    ) -> Dict[str, Any]:
        """Optimize resume content for a target job using LLM"""
        system_prompt = """You are an expert resume writer and ATS optimization specialist. 
Optimize resume content to match job descriptions while maintaining authenticity."""
        
        prompt = f"""Given this target job description:
{target_job}

Optimize the following resume components:

Experiences: {str(experiences)[:1000]}
Skills: {skills}
Education: {str(education)[:500]}
Portfolio: {str(portfolio_items)[:500] if portfolio_items else "None"}

Provide optimized versions that:
1. Include relevant keywords from the job description
2. Highlight matching experiences and skills
3. Maintain truthfulness and accuracy
4. Improve ATS compatibility

Return the optimized content in a structured format."""
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ]
        
        response = self.rag_engine.llm_client.chat_completion(messages=messages)
        
        # Parse response (simplified - in production, use structured output)
        # For now, return original data as LLM optimization would need more sophisticated parsing
        return {
            "experiences": experiences,
            "skills": skills,
        }
    
    def _generate_summary(
        self,
        experiences: List[Dict[str, Any]],
        skills: List[str],
        target_job: Optional[str],
    ) -> str:
        """Generate professional summary"""
        system_prompt = """You are an expert resume writer. Write concise, impactful professional summaries."""
        
        # Extract key information
        experience_summary = "\n".join([
            f"- {exp.get('title', 'N/A')} at {exp.get('company', 'N/A')}: {exp.get('description', '')[:100]}"
            for exp in experiences[:3]
        ])
        
        top_skills = ", ".join(skills[:10])
        
        prompt = f"""Write a professional summary (2-3 sentences) for a resume.

Key Experiences:
{experience_summary}

Top Skills: {top_skills}

Target Job: {target_job if target_job else "General"}
"""
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ]
        
        summary = self.rag_engine.llm_client.chat_completion(messages=messages)
        return summary.strip()
    
    def _create_docx_resume(
        self,
        resume_data: Dict[str, Any],
        user_id: Optional[str],
    ) -> str:
        """Create a DOCX resume document"""
        doc = Document()
        
        # Configure styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title (Name would go here - placeholder)
        title = doc.add_paragraph()
        title_run = title.add_run("PROFESSIONAL RESUME")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Professional Summary
        if resume_data.get("summary"):
            doc.add_paragraph("PROFESSIONAL SUMMARY", style='Heading 1')
            summary_para = doc.add_paragraph(resume_data["summary"])
            doc.add_paragraph()  # Spacing
        
        # Skills
        if resume_data.get("skills"):
            doc.add_paragraph("SKILLS", style='Heading 1')
            skills_text = " • ".join(resume_data["skills"])
            doc.add_paragraph(skills_text)
            doc.add_paragraph()  # Spacing
        
        # Experience
        if resume_data.get("experiences"):
            doc.add_paragraph("PROFESSIONAL EXPERIENCE", style='Heading 1')
            
            for exp in resume_data["experiences"]:
                # Job title and company
                job_header = doc.add_paragraph()
                job_title_run = job_header.add_run(exp.get("title", "N/A"))
                job_title_run.bold = True
                job_header.add_run(f" | {exp.get('company', 'N/A')}")
                
                # Duration
                if exp.get("duration"):
                    duration_para = doc.add_paragraph(exp["duration"])
                    duration_para.paragraph_format.left_indent = Inches(0.25)
                
                # Description
                if exp.get("description"):
                    desc_para = doc.add_paragraph(exp["description"])
                    desc_para.paragraph_format.left_indent = Inches(0.25)
                
                doc.add_paragraph()  # Spacing
        
        # Education
        if resume_data.get("education"):
            doc.add_paragraph("EDUCATION", style='Heading 1')
            
            for edu in resume_data["education"]:
                edu_para = doc.add_paragraph()
                degree_run = edu_para.add_run(edu.get("degree", "N/A"))
                degree_run.bold = True
                edu_para.add_run(f" | {edu.get('institution', 'N/A')}")
                
                if edu.get("year"):
                    year_para = doc.add_paragraph(edu["year"])
                    year_para.paragraph_format.left_indent = Inches(0.25)
                
                if edu.get("details"):
                    details_para = doc.add_paragraph(edu["details"])
                    details_para.paragraph_format.left_indent = Inches(0.25)
                
                doc.add_paragraph()  # Spacing
        
        # Portfolio/Projects (optional)
        if resume_data.get("portfolio"):
            doc.add_paragraph("PROJECTS", style='Heading 1')
            
            for project in resume_data["portfolio"]:
                proj_para = doc.add_paragraph()
                proj_name_run = proj_para.add_run(project.get("name", "N/A"))
                proj_name_run.bold = True
                
                if project.get("description"):
                    desc_para = doc.add_paragraph(project["description"])
                    desc_para.paragraph_format.left_indent = Inches(0.25)
                
                doc.add_paragraph()  # Spacing
        
        # Save document
        filename = "resume.docx"
        if user_id:
            output_path = self.file_handler.output_dir / user_id / filename
            output_path.parent.mkdir(parents=True, exist_ok=True)
        else:
            output_path = self.file_handler.output_dir / filename
        
        doc.save(str(output_path))
        return str(output_path)
    
    def _create_txt_resume(
        self,
        resume_data: Dict[str, Any],
        user_id: Optional[str],
    ) -> str:
        """Create a TXT resume document"""
        lines = []
        
        lines.append("=" * 60)
        lines.append("PROFESSIONAL RESUME")
        lines.append("=" * 60)
        lines.append("")
        
        # Summary
        if resume_data.get("summary"):
            lines.append("PROFESSIONAL SUMMARY")
            lines.append("-" * 60)
            lines.append(resume_data["summary"])
            lines.append("")
        
        # Skills
        if resume_data.get("skills"):
            lines.append("SKILLS")
            lines.append("-" * 60)
            lines.append(" • ".join(resume_data["skills"]))
            lines.append("")
        
        # Experience
        if resume_data.get("experiences"):
            lines.append("PROFESSIONAL EXPERIENCE")
            lines.append("-" * 60)
            
            for exp in resume_data["experiences"]:
                lines.append(f"{exp.get('title', 'N/A')} | {exp.get('company', 'N/A')}")
                if exp.get("duration"):
                    lines.append(f"  {exp['duration']}")
                if exp.get("description"):
                    lines.append(f"  {exp['description']}")
                lines.append("")
        
        # Education
        if resume_data.get("education"):
            lines.append("EDUCATION")
            lines.append("-" * 60)
            
            for edu in resume_data["education"]:
                lines.append(f"{edu.get('degree', 'N/A')} | {edu.get('institution', 'N/A')}")
                if edu.get("year"):
                    lines.append(f"  {edu['year']}")
                if edu.get("details"):
                    lines.append(f"  {edu['details']}")
                lines.append("")
        
        # Portfolio
        if resume_data.get("portfolio"):
            lines.append("PROJECTS")
            lines.append("-" * 60)
            
            for project in resume_data["portfolio"]:
                lines.append(project.get("name", "N/A"))
                if project.get("description"):
                    lines.append(f"  {project['description']}")
                lines.append("")
        
        resume_text = "\n".join(lines)
        
        # Save file
        filename = "resume.txt"
        file_path = self.file_handler.save_output(resume_text, filename, user_id)
        
        return file_path
    
    def _generate_resume_text(self, resume_data: Dict[str, Any]) -> str:
        """Generate plain text version of resume for analysis"""
        text_parts = []
        
        if resume_data.get("summary"):
            text_parts.append(resume_data["summary"])
        
        if resume_data.get("skills"):
            text_parts.append(" ".join(resume_data["skills"]))
        
        if resume_data.get("experiences"):
            for exp in resume_data["experiences"]:
                exp_text = f"{exp.get('title', '')} {exp.get('company', '')} {exp.get('description', '')}"
                text_parts.append(exp_text)
        
        if resume_data.get("education"):
            for edu in resume_data["education"]:
                edu_text = f"{edu.get('degree', '')} {edu.get('institution', '')} {edu.get('details', '')}"
                text_parts.append(edu_text)
        
        return "\n".join(text_parts)

