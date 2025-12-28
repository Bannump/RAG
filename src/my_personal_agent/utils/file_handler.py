"""
File handling utilities for documents (PDF, DOCX, TXT)
"""
import os
from pathlib import Path
from typing import Optional, List, Dict, Any
import PyPDF2
import pdfplumber
from docx import Document
from src.my_personal_agent.config import settings


class FileHandler:
    """Handle file operations for various document types"""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
    
    def __init__(self):
        self.uploads_dir = Path(settings.uploads_dir)
        self.output_dir = Path(settings.output_dir)
        self.uploads_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def save_upload(self, file_path: str, user_id: Optional[str] = None) -> str:
        """Save uploaded file and return new path"""
        source_path = Path(file_path)
        
        # Create user-specific directory if user_id provided
        if user_id:
            user_dir = self.uploads_dir / user_id
            user_dir.mkdir(parents=True, exist_ok=True)
            dest_path = user_dir / source_path.name
        else:
            dest_path = self.uploads_dir / source_path.name
        
        # Copy file
        import shutil
        shutil.copy2(source_path, dest_path)
        
        return str(dest_path)
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from document"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == ".pdf":
            return self._extract_pdf_text(file_path)
        elif extension == ".docx":
            return self._extract_docx_text(file_path)
        elif extension in {".txt", ".md"}:
            return self._read_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using pdfplumber (more accurate)"""
        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception:
            # Fallback to PyPDF2
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        
        return "\n".join(text_parts)
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    def _read_text_file(self, file_path: Path) -> str:
        """Read text from plain text file"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    
    def save_output(self, content: str, filename: str, user_id: Optional[str] = None) -> str:
        """Save output file"""
        if user_id:
            user_dir = self.output_dir / user_id
            user_dir.mkdir(parents=True, exist_ok=True)
            output_path = user_dir / filename
        else:
            output_path = self.output_dir / filename
        
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(content)
        
        return str(output_path)
    
    def save_docx(self, paragraphs: List[Dict[str, Any]], filename: str, user_id: Optional[str] = None) -> str:
        """Save content as DOCX file"""
        doc = Document()
        
        for para_data in paragraphs:
            para = doc.add_paragraph(para_data.get("text", ""))
            
            # Apply formatting if specified
            if para_data.get("bold"):
                for run in para.runs:
                    run.bold = True
            if para_data.get("italic"):
                for run in para.runs:
                    run.italic = True
        
        if user_id:
            user_dir = self.output_dir / user_id
            user_dir.mkdir(parents=True, exist_ok=True)
            output_path = user_dir / filename
        else:
            output_path = self.output_dir / filename
        
        doc.save(output_path)
        return str(output_path)

